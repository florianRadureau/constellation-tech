"""
Service d'extraction de texte depuis CV (PDF et DOCX).
"""
import io
import logging
import re

from docx import Document
from pypdf import PdfReader

from exceptions import CVParserError

logger = logging.getLogger(__name__)


class CVParser:
    """Parse les CV pour en extraire le texte brut."""

    # Constantes de classe
    SUPPORTED_EXTENSIONS = ['pdf', 'docx', 'doc']
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB en bytes

    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        """
        Point d'entrée principal : extrait le texte d'un CV.

        Args:
            file_content: Contenu binaire du fichier
            filename: Nom du fichier (pour détecter l'extension)

        Returns:
            Texte extrait du CV

        Raises:
            CVParserError: Si l'extraction échoue ou l'extension n'est pas bonne
        """
        try:
            CVParser.validate_file(file_content, filename)
        except CVParserError:
            # Re-raise nos propres erreurs
            raise

        file_extension = CVParser.get_file_extension(filename)
        if (file_extension == "pdf"):
            return CVParser.extract_text_from_pdf(file_content)
        elif (file_extension == "docx" or file_extension == "doc"):
            return CVParser.extract_text_from_docx(file_content)


    @staticmethod
    def get_file_extension(filename: str) -> str:
        """"
        Extrait l'extension du fichier

        Args:
            filename : Nom du fichier

        CVParserError: Si le nom du fichier n'est pas conforme

        Returns:
            Extension du fichier
        """
        try:
            extension = filename.lower().split(".")[-1]
        except Exception as e:
            raise CVParserError(f"Le nom du fichier est incorrect : {str(e)}")
        return extension



    @staticmethod
    def validate_file(file_content: bytes, filename: str) -> None:
        """
        Valide qu'un fichier peut être traité.

        Args:
            file_content: Contenu du fichier
            filename: Nom du fichier

        Raises:
            CVParserError: Si la validation échoue
        """
        # Validation 1 : Extension
        extension = CVParser.get_file_extension(filename)

        if extension not in CVParser.SUPPORTED_EXTENSIONS:
            raise CVParserError(
                f"Format de fichier non supporté : .{extension}. "
                f"Formats acceptés : {', '.join(CVParser.SUPPORTED_EXTENSIONS).upper()}"
            )

        # Validation 2 : Taille
        file_size = len(file_content)
        if file_size > CVParser.MAX_FILE_SIZE:
            size_mb = file_size / (1024 * 1024)
            max_mb = CVParser.MAX_FILE_SIZE / (1024 * 1024)
            raise CVParserError(
                f"Fichier trop volumineux ({size_mb:.1f} MB). "
                f"Taille maximale : {max_mb:.0f} MB"
            )

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Nettoie le texte extrait.

        Opérations :
        - Remplace sauts de ligne multiples par un seul
        - Remplace espaces/tabs multiples par un seul espace
        - Enlève espaces début/fin de chaque ligne
        - Trim général

        Args:
            text: Texte brut extrait

        Returns:
            Texte nettoyé
        """
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)

        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        text = '\n'.join(lines)
        text = text.strip()

        return text

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extrait le texte d'un PDF.

        Tente d'abord l'extraction normale (PDFs avec texte).
        Si aucun texte extractible (PDF scanné), utilise OCR automatiquement.

        Args:
            file_content: Contenu binaire du fichier PDF

        Returns:
            Texte extrait (nettoyé)

        Raises:
            CVParserError: Si l'extraction échoue
        """
        try:
            # 1. Tentative extraction normale (PDFs avec texte)
            with io.BytesIO(file_content) as pdf_file:
                reader = PdfReader(pdf_file)

                # Vérifier qu'il y a des pages
                if len(reader.pages) == 0:
                    raise CVParserError("Le PDF ne contient aucune page")

                # Extraire le texte de chaque page
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                # Vérifier qu'on a du texte
                if not text_parts:
                    raise CVParserError("Aucun texte extractible")

                # Joindre et nettoyer
                full_text = "\n".join(text_parts)
                return CVParser._clean_text(full_text)

        except CVParserError as e:
            # Si "Aucun texte extractible", essayer OCR en fallback
            if "Aucun texte extractible" in str(e):
                logger.info(
                    "PDF scanné détecté (pas de texte extractible), "
                    "utilisation OCR automatique"
                )
                try:
                    from services.ocr_service import OCRService

                    ocr_service = OCRService()
                    text = ocr_service.extract_text_from_pdf(file_content)
                    logger.info("OCR extraction successful")
                    return text
                except Exception as ocr_error:
                    logger.error(f"OCR extraction failed: {ocr_error}", exc_info=True)
                    raise CVParserError(
                        f"Impossible d'extraire le texte (PDF scanné, OCR échoué): {str(ocr_error)}"
                    ) from ocr_error
            # Autres erreurs CVParserError: re-raise
            raise
        except Exception as e:
            # Capture les autres erreurs
            logger.error(f"PDF extraction failed: {e}", exc_info=True)
            raise CVParserError(f"Erreur lecture PDF: {str(e)}") from e

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extrait le texte d'un DOCX.

        Args:
            file_content: Contenu binaire du fichier DOCX

        Returns:
            Texte extrait (nettoyé)

        Raises:
            CVParserError: Si l'extraction échoue
        """
        try:
            with io.BytesIO(file_content) as docx_file:
                doc = Document(docx_file)

                text_parts = []

                # Extraire des paragraphes
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # Extraire des tableaux
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)

                # Vérifier qu'on a du texte
                if not text_parts:
                    raise CVParserError("Le document DOCX est vide")

                # Joindre et nettoyer
                full_text = "\n".join(text_parts)
                return CVParser._clean_text(full_text)

        except CVParserError:
            raise
        except Exception as e:
            raise CVParserError(f"Erreur lors de la lecture du DOCX: {str(e)}")

